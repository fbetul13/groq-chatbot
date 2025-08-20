#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Debug Markdown to Word Converter
"""

import os
from pathlib import Path

def convert_md_to_word():
    """Markdown dosyasını Word belgesine dönüştürür."""
    
    # Gerekli kütüphaneleri kontrol et ve yükle
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        print("✅ python-docx kütüphanesi yüklendi")
    except ImportError:
        print("❌ python-docx kütüphanesi bulunamadı")
        return False
    
    # Markdown dosyasının yolunu belirle
    md_file = Path("/Users/betul/Desktop/bitirme_projesi_taslak.md")
    
    if not md_file.exists():
        print(f"❌ {md_file} dosyası bulunamadı!")
        return False
    
    # Markdown içeriğini oku
    try:
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
        print(f"✅ Markdown dosyası okundu: {len(md_content)} karakter")
    except Exception as e:
        print(f"❌ Markdown dosyası okunamadı: {e}")
        return False
    
    # Word belgesi oluştur
    doc = Document()
    print("✅ Word belgesi oluşturuldu")
    
    # Test paragrafı ekle
    doc.add_paragraph("TEST PARAGRAFI - Bu çalışıyorsa script doğru")
    print("✅ Test paragrafı eklendi")
    
    # Markdown içeriğini satır satır işle
    lines = md_content.split('\n')
    print(f"📝 {len(lines)} satır işleniyor...")
    
    paragraph_count = 0
    
    for i, line in enumerate(lines):
        line = line.strip()
        
        if not line:
            # Boş satır
            doc.add_paragraph()
            paragraph_count += 1
            continue
        
        # Başlık kontrolü
        if line.startswith('## '):
            # Ana başlık
            title = line[3:].strip()
            print(f"📋 Ana başlık: {title}")
            if title == "Kapak Bilgileri":
                # Kapak sayfası için özel format
                doc.add_page_break()
                title_para = doc.add_paragraph(title)
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_para.style = doc.styles['Title']
                doc.add_paragraph()  # Boşluk
                paragraph_count += 3
            else:
                doc.add_heading(title, level=1)
                paragraph_count += 1
            continue
        
        elif line.startswith('### '):
            # Alt başlık
            title = line[4:].strip()
            print(f"📋 Alt başlık: {title}")
            doc.add_heading(title, level=2)
            paragraph_count += 1
            continue
        
        elif line.startswith('#### '):
            # Alt-alt başlık
            title = line[5:].strip()
            print(f"📋 Alt-alt başlık: {title}")
            doc.add_heading(title, level=3)
            paragraph_count += 1
            continue
        
        elif line.startswith('---'):
            # Ayırıcı çizgi
            doc.add_paragraph("─" * 50)
            paragraph_count += 1
            continue
        
        elif line.startswith('```'):
            # Kod bloğu - basit işleme
            if line == '```':
                # Kod bloğu başlangıcı
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                
                if code_lines:
                    # Kod bloğunu ekle
                    code_para = doc.add_paragraph()
                    code_para.style = 'No Spacing'
                    for code_line in code_lines:
                        code_para.add_run(code_line + '\n')
                    paragraph_count += 1
                
                if i < len(lines) and lines[i].strip().startswith('```'):
                    i += 1  # Kapanış ``` işaretini atla
                continue
            else:
                # Tek satır kod
                code_text = line[3:].strip()
                if code_text:
                    code_para = doc.add_paragraph(code_text)
                    code_para.style = 'No Spacing'
                    paragraph_count += 1
                continue
        
        elif line.startswith('- **'):
            # Liste öğesi
            content = line[2:].strip()
            if content.startswith('**') and '**:' in content:
                # Kalın başlık + açıklama
                parts = content.split('**:', 1)
                if len(parts) == 2:
                    title, desc = parts
                    title = title[2:]  # ** işaretlerini kaldır
                    para = doc.add_paragraph()
                    para.add_run(title).bold = True
                    para.add_run(': ' + desc)
                    paragraph_count += 1
            else:
                # Normal liste öğesi
                doc.add_paragraph(content, style='List Bullet')
                paragraph_count += 1
            continue
        
        elif line.startswith('|'):
            # Tablo - basit işleme
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            if table_lines:
                # Basit tablo oluştur
                try:
                    create_simple_table(doc, table_lines)
                    paragraph_count += 1
                except Exception as e:
                    print(f"⚠️ Tablo oluşturulamadı: {e}")
                    doc.add_paragraph(f"[Tablo: {len(table_lines)} satır]")
                    paragraph_count += 1
            continue
        
        elif line.startswith('```mermaid'):
            # Mermaid diyagramı (şimdilik atlanıyor)
            while i < len(lines) and not lines[i].strip().startswith('```'):
                i += 1
            if i < len(lines):
                i += 1  # Kapanış ``` işaretini atla
            doc.add_paragraph("[Mermaid Diyagramı - Word'de görüntülenemez]")
            paragraph_count += 1
            continue
        
        else:
            # Normal metin
            if line.startswith('[') and line.endswith(']'):
                # Placeholder metni
                para = doc.add_paragraph(line)
                para.runs[0].font.italic = True
                paragraph_count += 1
            else:
                doc.add_paragraph(line)
                paragraph_count += 1
            continue
    
    print(f"📄 Word belgesi oluşturuldu, {paragraph_count} paragraf")
    print(f"📄 Gerçek paragraf sayısı: {len(doc.paragraphs)}")
    
    # Word dosyasını kaydet
    output_file = "bitirme_projesi_taslak_debug.docx"
    try:
        doc.save(output_file)
        print(f"✅ Word belgesi başarıyla kaydedildi: {output_file}")
        print(f"📁 Dosya konumu: {os.path.abspath(output_file)}")
        return True
    except Exception as e:
        print(f"❌ Word dosyası kaydedilemedi: {e}")
        return False

def create_simple_table(doc, table_lines):
    """Basit tablo oluşturur."""
    if len(table_lines) < 2:
        return
    
    try:
        # İlk satırı başlık olarak kullan
        headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
        
        # Tablo oluştur
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Başlık satırını ekle
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            if header_cells[i].paragraphs[0].runs:
                header_cells[i].paragraphs[0].runs[0].bold = True
        
        # Veri satırlarını ekle (--- satırını atla)
        for line in table_lines[2:]:
            if line.strip() and not line.strip().startswith('|'):
                continue
            
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if len(cells) == len(headers):
                row_cells = table.add_row().cells
                for i, cell_text in enumerate(cells):
                    row_cells[i].text = cell_text
    except Exception as e:
        print(f"⚠️ Tablo oluşturma hatası: {e}")
        doc.add_paragraph(f"[Tablo hatası: {e}]")

def main():
    """Ana fonksiyon."""
    print("🔄 Debug: Markdown dosyası Word belgesine dönüştürülüyor...")
    print("=" * 60)
    
    success = convert_md_to_word()
    
    if success:
        print("\n🎉 Dönüştürme tamamlandı!")
        print("\n📋 Sonraki adımlar:")
        print("1. Word dosyasını açın")
        print("2. İçeriği kontrol edin")
        print("3. Debug çıktılarını inceleyin")
    else:
        print("\n❌ Dönüştürme başarısız!")
        print("Lütfen hata mesajlarını kontrol edin.")

if __name__ == "__main__":
    main() 
#!/usr/bin/env python3
from docx import Document

# Basit Word belgesi oluştur
doc = Document()

# Başlık ekle
doc.add_heading('Test Başlık', 0)

# Paragraf ekle
doc.add_paragraph('Bu bir test paragrafıdır.')

# Liste ekle
doc.add_paragraph('Liste öğesi 1', style='List Bullet')
doc.add_paragraph('Liste öğesi 2', style='List Bullet')

# Kaydet
doc.save('test.docx')
print("Test Word dosyası oluşturuldu: test.docx") 
#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Basit Markdown to Word Converter
Bu script, bitirme_projesi_taslak.md dosyasını Word belgesine dönüştürür.
"""

import os
from pathlib import Path

def convert_md_to_word():
    """Markdown dosyasını Word belgesine dönüştürür."""
    
    # Gerekli kütüphaneleri kontrol et ve yükle
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("Gerekli kütüphaneler yükleniyor...")
        os.system("pip3 install python-docx")
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            print("Hata: python-docx kütüphanesi yüklenemedi.")
            return False
    
    # Markdown dosyasının yolunu belirle
    md_file = Path("/Users/betul/Desktop/bitirme_projesi_taslak.md")
    
    if not md_file.exists():
        print(f"Hata: {md_file} dosyası bulunamadı!")
        return False
    
    # Markdown içeriğini oku
    try:
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
        print(f"✅ Markdown dosyası okundu: {len(md_content)} karakter")
    except Exception as e:
        print(f"Markdown dosyası okunamadı: {e}")
        return False
    
    # Word belgesi oluştur
    doc = Document()
    
    # Başlık stillerini ayarla
    title_style = doc.styles['Title']
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    
    heading1_style = doc.styles['Heading 1']
    heading1_style.font.size = Pt(16)
    heading1_style.font.bold = True
    
    heading2_style = doc.styles['Heading 2']
    heading2_style.font.size = Pt(14)
    heading2_style.font.bold = True
    
    heading3_style = doc.styles['Heading 3']
    heading3_style.font.size = Pt(12)
    heading3_style.font.bold = True
    
    # Markdown içeriğini satır satır işle
    lines = md_content.split('\n')
    print(f"📝 {len(lines)} satır işleniyor...")
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            # Boş satır
            doc.add_paragraph()
            i += 1
            continue
        
        # Başlık kontrolü
        if line.startswith('## '):
            # Ana başlık
            title = line[3:].strip()
            print(f"📋 Ana başlık: {title}")
            if title == "Kapak Bilgileri":
                # Kapak sayfası için özel format
                doc.add_page_break()
                title_para = doc.add_paragraph(title)
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_para.style = title_style
                doc.add_paragraph()  # Boşluk
            else:
                doc.add_heading(title, level=1)
            i += 1
            continue
        
        elif line.startswith('### '):
            # Alt başlık
            title = line[4:].strip()
            print(f"📋 Alt başlık: {title}")
            doc.add_heading(title, level=2)
            i += 1
            continue
        
        elif line.startswith('#### '):
            # Alt-alt başlık
            title = line[5:].strip()
            print(f"📋 Alt-alt başlık: {title}")
            doc.add_heading(title, level=3)
            i += 1
            continue
        
        elif line.startswith('---'):
            # Ayırıcı çizgi
            doc.add_paragraph("─" * 50)
            i += 1
            continue
        
        elif line.startswith('```'):
            # Kod bloğu
            if line == '```':
                # Kod bloğu başlangıcı
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                
                if code_lines:
                    # Kod bloğunu ekle
                    code_para = doc.add_paragraph()
                    code_para.style = 'No Spacing'
                    for code_line in code_lines:
                        code_para.add_run(code_line + '\n')
                
                if i < len(lines) and lines[i].strip().startswith('```'):
                    i += 1  # Kapanış ``` işaretini atla
                continue
            else:
                # Tek satır kod
                code_text = line[3:].strip()
                if code_text:
                    code_para = doc.add_paragraph(code_text)
                    code_para.style = 'No Spacing'
                i += 1
                continue
        
        elif line.startswith('- **'):
            # Liste öğesi
            content = line[2:].strip()
            if content.startswith('**') and '**:' in content:
                # Kalın başlık + açıklama
                parts = content.split('**:', 1)
                if len(parts) == 2:
                    title, desc = parts
                    title = title[2:]  # ** işaretlerini kaldır
                    para = doc.add_paragraph()
                    para.add_run(title).bold = True
                    para.add_run(': ' + desc)
            else:
                # Normal liste öğesi
                doc.add_paragraph(content, style='List Bullet')
            i += 1
            continue
        
        elif line.startswith('|'):
            # Tablo
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            if table_lines:
                # Tablo oluştur
                create_table_from_markdown(doc, table_lines)
            continue
        
        elif line.startswith('```mermaid'):
            # Mermaid diyagramı (şimdilik atlanıyor)
            while i < len(lines) and not lines[i].strip().startswith('```'):
                i += 1
            if i < len(lines):
                i += 1  # Kapanış ``` işaretini atla
            doc.add_paragraph("[Mermaid Diyagramı - Word'de görüntülenemez]")
            continue
        
        else:
            # Normal metin
            if line.startswith('[') and line.endswith(']'):
                # Placeholder metni
                para = doc.add_paragraph(line)
                para.runs[0].font.italic = True
            else:
                doc.add_paragraph(line)
            i += 1
            continue
    
    print(f"📄 Word belgesi oluşturuldu, {len(doc.paragraphs)} paragraf")
    
    # Word dosyasını kaydet
    output_file = "bitirme_projesi_taslak.docx"
    try:
        doc.save(output_file)
        print(f"✅ Word belgesi başarıyla kaydedildi: {output_file}")
        print(f"📁 Dosya konumu: {os.path.abspath(output_file)}")
        return True
    except Exception as e:
        print(f"❌ Word dosyası kaydedilemedi: {e}")
        return False

def create_table_from_markdown(doc, table_lines):
    """Markdown tablo formatından Word tablosu oluşturur."""
    if len(table_lines) < 2:
        return
    
    # İlk satırı başlık olarak kullan
    headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
    
    # Tablo oluştur
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Başlık satırını ekle
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
    
    # Veri satırlarını ekle (--- satırını atla)
    for line in table_lines[2:]:
        if line.strip() and not line.strip().startswith('|'):
            continue
        
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if len(cells) == len(headers):
            row_cells = table.add_row().cells
            for i, cell_text in enumerate(cells):
                row_cells[i].text = cell_text

def main():
    """Ana fonksiyon."""
    print("🔄 Markdown dosyası Word belgesine dönüştürülüyor...")
    print("=" * 50)
    
    success = convert_md_to_word()
    
    if success:
        print("\n🎉 Dönüştürme tamamlandı!")
        print("\n📋 Sonraki adımlar:")
        print("1. Word dosyasını açın")
        print("2. Gerekli bilgileri doldurun")
        print("3. Formatlamayı kontrol edin")
        print("4. Gerekirse düzenlemeler yapın")
    else:
        print("\n❌ Dönüştürme başarısız!")
        print("Lütfen hata mesajlarını kontrol edin.")

if __name__ == "__main__":
    main() 